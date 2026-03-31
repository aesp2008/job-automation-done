import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from backend.app.services.skill_lexicon import skills_from_known_lexicon

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")


def _emails_from_text(text: str) -> list[str]:
    found = _EMAIL_RE.findall(text)
    # De-dupe while preserving order
    seen: set[str] = set()
    out: list[str] = []
    for e in found:
        key = e.lower()
        if key not in seen:
            seen.add(key)
            out.append(e)
        if len(out) >= 5:
            break
    return out


def _extract_text(filename: str, content: bytes) -> tuple[str, str | None]:
    """Return (text, optional_extraction_note)."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        try:
            reader = PdfReader(BytesIO(content))
            parts: list[str] = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            text = "\n".join(parts)
            if not text.strip():
                return (
                    "",
                    "PDF had no extractable text (it may be a scan or image-only).",
                )
            return text, None
        except PdfReadError:
            return "", "Could not read PDF (file may be corrupted or encrypted)."
        except Exception:
            return "", "Could not read PDF."

    if ext == ".docx":
        try:
            doc = Document(BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text)
            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + "\t".join(cell.text for cell in row.cells)
            if not text.strip():
                return "", "DOCX contained no readable paragraphs."
            return text, None
        except Exception:
            return "", "Could not read DOCX."

    if ext == ".doc":
        return (
            "",
            "Legacy .doc is not supported. Please export as PDF or DOCX.",
        )

    raw = content.decode("utf-8", errors="ignore")
    return raw, None


def extract_resume_plain_text(filename: str, content: bytes) -> tuple[str, str | None]:
    """Public helper for tailoring and other services that need raw resume text."""
    return _extract_text(filename, content)


def parse_resume_bytes(filename: str, content: bytes) -> dict:
    """Extract text from PDF / DOCX / plain text and return a structured summary."""
    text, extraction_note = _extract_text(filename, content)
    skills = skills_from_known_lexicon(text, limit=16)
    emails = _emails_from_text(text)

    summary_parts = [
        "Keyword and contact extraction from resume text.",
    ]
    if extraction_note:
        summary_parts.insert(0, extraction_note)

    return {
        "filename": Path(filename).name,
        "file_size_kb": round(len(content) / 1024, 2),
        "extension": Path(filename).suffix.lower(),
        "skills_detected": skills,
        "emails_found": emails,
        "text_preview": (text[:400] + "…") if len(text) > 400 else text,
        "summary": " ".join(summary_parts),
    }
