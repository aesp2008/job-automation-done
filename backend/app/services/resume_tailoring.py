"""JD-aware resume draft: reorder skills and suggest bullets without external LLM."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from backend.app.services import resume_parser
from backend.app.services.skill_lexicon import skills_from_known_lexicon


def load_user_resume_plain_text(resume_path: str | None) -> tuple[str, str | None]:
    """Return (text, error_note)."""
    if not resume_path:
        return "", "No resume on file. Upload one under Settings → Preferences."
    path = Path(resume_path)
    if not path.is_file():
        return "", "Resume file was not found on the server (path missing)."
    raw = path.read_bytes()
    text, note = resume_parser.extract_resume_plain_text(path.name, raw)
    if note and not text.strip():
        return "", note
    return text, note


def build_tailoring_payload(
    resume_plain_text: str,
    *,
    job_title: str,
    company: str,
    jd_description: str,
    resume_extraction_note: str | None = None,
) -> dict:
    jd_skills = skills_from_known_lexicon(jd_description or "", limit=24)
    resume_skills = skills_from_known_lexicon(resume_plain_text or "", limit=24)
    resume_set = set(resume_skills)
    matched = [s for s in jd_skills if s in resume_set]
    gaps = [s for s in jd_skills if s not in resume_set]

    ordered_skills: list[str] = []
    for s in jd_skills:
        if s not in ordered_skills:
            ordered_skills.append(s)
    for s in resume_skills:
        if s not in ordered_skills:
            ordered_skills.append(s)

    bullets: list[str] = []
    for s in matched[:5]:
        bullets.append(
            f"Shipped features using {s} in production environments aligned with this role."
        )
    if not bullets and resume_skills:
        bullets.append(
            f"Strong delivery background across {', '.join(resume_skills[:5])}."
        )
    for s in gaps[:4]:
        bullets.append(
            f"Keen to expand hands-on depth with {s}, complementing adjacent stack experience."
        )

    if matched:
        summary = (
            f"Focus for {job_title} at {company}: emphasize "
            + ", ".join(matched[:8])
            + " as called out in the posting."
        )
    else:
        summary = (
            f"Focus for {job_title} at {company}: align narrative with posting themes; "
            "highlight transferable depth from your stack."
        )

    notes_block = ""
    if resume_extraction_note:
        notes_block = f"\nNote on source resume: {resume_extraction_note}\n"

    full_text = f"""TAILORED RESUME DRAFT — {job_title} @ {company}
(Use as a guide; merge accurate facts from your real experience.)

# Professional summary
{summary}

# Skills (JD-relevant order)
{", ".join(ordered_skills) if ordered_skills else "—"}

# Suggested bullets (edit with real metrics & ownership)
{chr(10).join(f"- {b}" for b in bullets) if bullets else "—"}
{notes_block}
# Source resume (excerpt)
---
{resume_plain_text[:12_000]}
"""

    return {
        "job_title": job_title,
        "company": company,
        "jd_skills_highlighted": jd_skills,
        "resume_skills_detected": resume_skills,
        "skills_matched_with_jd": matched,
        "skills_gaps_vs_jd": gaps,
        "skills_section_ordered": ordered_skills,
        "suggested_bullets": bullets,
        "professional_summary": summary,
        "full_text_draft": full_text.strip(),
        "resume_extraction_note": resume_extraction_note,
        "resume_excerpt": (resume_plain_text or "")[:12_000],
    }


def tailoring_payload_to_docx_bytes(payload: dict) -> bytes:
    """Build a minimal .docx users can edit in Word from the tailoring payload."""
    doc = Document()
    doc.add_heading(f"Tailored draft — {payload['job_title']}", level=0)
    doc.add_paragraph(payload["company"])
    doc.add_heading("Professional summary", level=1)
    doc.add_paragraph(payload["professional_summary"])

    doc.add_heading("Skills (JD-relevant order)", level=1)
    doc.add_paragraph(", ".join(payload["skills_section_ordered"]) or "—")

    doc.add_heading("JD alignment", level=1)
    matched = payload["skills_matched_with_jd"]
    gaps = payload["skills_gaps_vs_jd"]
    doc.add_paragraph(
        "Matched: " + (", ".join(matched) if matched else "None listed."),
    )
    doc.add_paragraph(
        "Gaps to address: " + (", ".join(gaps) if gaps else "None listed."),
    )

    doc.add_heading("Suggested bullets (edit with your metrics)", level=1)
    for b in payload["suggested_bullets"]:
        doc.add_paragraph(b, style="List Bullet")

    if payload.get("resume_extraction_note"):
        doc.add_heading("Note on uploaded file", level=1)
        doc.add_paragraph(str(payload["resume_extraction_note"]))

    excerpt = (payload.get("resume_excerpt") or "").strip()
    if excerpt:
        doc.add_heading("Your resume text (excerpt)", level=1)
        doc.add_paragraph(excerpt)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
